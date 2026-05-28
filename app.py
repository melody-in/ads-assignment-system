import os
import sys
import json
import uuid
import subprocess
import shutil
from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
import io
import time
import copy
import re

app = Flask(__name__)
CORS(app)

# ─── Configuration ───────────────────────────────────────────────────────────

BASE_DIR = os.path.dirname(__file__)
DOCUMENTS_DIR = os.path.join(BASE_DIR, 'documents')
SETTINGS_PATH = os.path.join(DOCUMENTS_DIR, 'settings.json')
OUTPUT_DIR = os.path.join(BASE_DIR, 'generated')
_cleanup_done = False

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(DOCUMENTS_DIR, exist_ok=True)


def load_settings():
    """Load settings from documents/settings.json, creating defaults if missing."""
    defaults = {
        "template": "",
        "findName": "N. Akshit Vinay",
        "findPid": "25MSRSGIS001",
        "findCourse": "M.Sc.  Remote Sensing & Gis"
    }

    if not os.path.exists(SETTINGS_PATH):
        with open(SETTINGS_PATH, 'w') as f:
            json.dump(defaults, f, indent=2)
        print(f"[INFO] Created default settings at {SETTINGS_PATH}", file=sys.stderr)
        return defaults

    try:
        with open(SETTINGS_PATH, 'r') as f:
            settings = json.load(f)
            for key in defaults:
                if key not in settings:
                    settings[key] = defaults[key]
            return settings
    except Exception as e:
        print(f"[WARN] Could not read settings: {e}", file=sys.stderr)
        return defaults


def find_template():
    """Find the DOCX template in the documents/ folder."""
    settings = load_settings()
    template_name = settings.get('template', '')

    # If template is specified in settings, use it
    if template_name:
        path = os.path.join(DOCUMENTS_DIR, template_name)
        if os.path.exists(path) and not template_name.startswith('~$'):
            return path
        else:
            # Template in settings no longer exists or is a lock file, reset it
            print(f"[WARN] Template '{template_name}' not found or is lock file, searching for alternatives...", file=sys.stderr)
            settings['template'] = ''

    # Otherwise, look for any .docx file in the documents/ folder
    docx_files = [f for f in os.listdir(DOCUMENTS_DIR)
                  if f.lower().endswith('.docx') and not f.startswith('~$') and f != 'settings.json']
    if docx_files:
        docx_files.sort(key=lambda f: os.path.getmtime(os.path.join(DOCUMENTS_DIR, f)), reverse=True)
        chosen = docx_files[0]
        settings['template'] = chosen
        try:
            with open(SETTINGS_PATH, 'w') as f:
                json.dump(settings, f, indent=2)
        except Exception:
            pass
        return os.path.join(DOCUMENTS_DIR, chosen)

    # Fallback: look in project root
    root_docx = [f for f in os.listdir(BASE_DIR) if f.lower().endswith('.docx')]
    if root_docx:
        return os.path.join(BASE_DIR, root_docx[0])

    return None


# ─── Robust Run-Aware Find-and-Replace ────────────────────────────────────────

def _replace_in_paragraph(paragraph, find_text, replace_text):
    """
    Replace text across runs in a paragraph, preserving formatting.
    Handles cases where the target text is split across multiple runs.
    """
    full_text = paragraph.text
    if find_text not in full_text:
        return False

    # Build a map of character positions to runs
    runs = paragraph.runs
    if not runs:
        return False

    # Try simple run-by-run replacement first
    for run in runs:
        if find_text in run.text:
            run.text = run.text.replace(find_text, replace_text)
            return True

    # If not found in a single run, the text is split across runs.
    # We need to find which runs contain the target text and merge/replace.
    char_positions = []  # (run_index, char_index_in_run)
    for ri, run in enumerate(runs):
        for ci, ch in enumerate(run.text):
            char_positions.append((ri, ci))

    # Find the start position of find_text in the full text
    start_pos = full_text.find(find_text)
    if start_pos < 0 or start_pos + len(find_text) > len(char_positions):
        return False

    end_pos = start_pos + len(find_text)

    # Determine which runs are affected
    start_run_idx = char_positions[start_pos][0]
    end_run_idx = char_positions[end_pos - 1][0]
    start_char = char_positions[start_pos][1]
    end_char = char_positions[end_pos - 1][1]

    # Put the replacement text in the first affected run, preserving its formatting
    first_run = runs[start_run_idx]
    before = first_run.text[:start_char]
    first_run.text = before + replace_text

    # If the text spans to the last affected run, keep the tail
    if start_run_idx == end_run_idx:
        first_run.text += runs[end_run_idx].text[end_char + 1:]
    else:
        last_run = runs[end_run_idx]
        after = last_run.text[end_char + 1:]
        first_run.text += after

        # Clear intermediate and last runs
        for ri in range(start_run_idx + 1, end_run_idx + 1):
            runs[ri].text = ""

    return True


# ─── DOCX Modification ───────────────────────────────────────────────────────

def modify_docx_template(student_name, student_pid, student_course, output_path=None):
    """
    Modify the DOCX template by replacing ONLY the Submitted By fields.
    Preserves ALL original formatting.
    """
    template_path = find_template()
    if not template_path:
        raise FileNotFoundError("No DOCX template found. Place a .docx file in the 'documents/' folder.")

    settings = load_settings()
    find_name = settings.get('findName', 'N. Akshit Vinay')
    find_pid = settings.get('findPid', '25MSRSGIS001')
    find_course = settings.get('findCourse', 'M.Sc.  Remote Sensing & Gis')

    doc = Document(template_path)

    # ── Replace in all paragraphs ──
    for para in doc.paragraphs:
        _replace_in_paragraph(para, find_name, student_name)
        _replace_in_paragraph(para, find_pid, student_pid)
        _replace_in_paragraph(para, find_course, student_course)
        _replace_in_paragraph(para, 'M.Sc. GIS & Remote Sensing', student_course)
        _replace_in_paragraph(para, 'M.Sc.  Remote Sensing & Gis', student_course)

    # ── Replace in all tables ──
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, find_name, student_name)
                    _replace_in_paragraph(para, find_pid, student_pid)
                    _replace_in_paragraph(para, find_course, student_course)
                    _replace_in_paragraph(para, 'M.Sc. GIS & Remote Sensing', student_course)
                    _replace_in_paragraph(para, 'M.Sc.  Remote Sensing & Gis', student_course)

    # ── Replace in headers and footers ──
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header and header.is_linked_to_previous is False:
                for para in header.paragraphs:
                    _replace_in_paragraph(para, find_name, student_name)
                    _replace_in_paragraph(para, find_pid, student_pid)
                    _replace_in_paragraph(para, find_course, student_course)
                    _replace_in_paragraph(para, 'M.Sc. GIS & Remote Sensing', student_course)
                    _replace_in_paragraph(para, 'M.Sc.  Remote Sensing & Gis', student_course)
        for footer in [section.footer, section.first_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                for para in footer.paragraphs:
                    _replace_in_paragraph(para, find_name, student_name)
                    _replace_in_paragraph(para, find_pid, student_pid)
                    _replace_in_paragraph(para, find_course, student_course)
                    _replace_in_paragraph(para, 'M.Sc. GIS & Remote Sensing', student_course)
                    _replace_in_paragraph(para, 'M.Sc.  Remote Sensing & Gis', student_course)

    if output_path:
        doc.save(output_path)
        return output_path

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── PDF Generation via LibreOffice ──────────────────────────────────────────

def convert_docx_to_pdf(docx_path, output_dir):
    """
    Convert DOCX to PDF using LibreOffice headless.
    Tries 'libreoffice' first, then 'soffice' as fallback.
    This preserves ALL original formatting exactly.
    """
    # Try both common binary names for LibreOffice
    for binary in ['libreoffice', 'soffice']:
        try:
            print(f"[PDF] Trying {binary} to convert {os.path.basename(docx_path)}...", file=sys.stderr)
            result = subprocess.run(
                [
                    binary, '--headless', '--convert-to', 'pdf',
                    '--outdir', output_dir, docx_path
                ],
                capture_output=True, text=True, timeout=120
            )
            if result.returncode != 0:
                print(f"[PDF] {binary} error (code {result.returncode}): {result.stderr}", file=sys.stderr)
                continue

            # LibreOffice outputs the PDF with the same base name
            base_name = os.path.splitext(os.path.basename(docx_path))[0]
            pdf_path = os.path.join(output_dir, base_name + '.pdf')
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print(f"[PDF] Success: {pdf_path} ({os.path.getsize(pdf_path)} bytes)", file=sys.stderr)
                return pdf_path
            else:
                print(f"[PDF] {binary} ran but PDF file not found or empty at {pdf_path}", file=sys.stderr)
                continue

        except FileNotFoundError:
            print(f"[PDF] {binary} not found, trying next...", file=sys.stderr)
            continue
        except subprocess.TimeoutExpired:
            print(f"[PDF] {binary} conversion timed out.", file=sys.stderr)
            continue
        except Exception as e:
            print(f"[PDF] {binary} error: {e}", file=sys.stderr)
            continue

    print("[PDF] All conversion methods failed. Neither libreoffice nor soffice available.", file=sys.stderr)
    return None


# ─── Routes ──────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    global _cleanup_done
    if not _cleanup_done:
        _cleanup_old_files()
        tmpl = find_template()
        if tmpl:
            print(f"[OK] Template found: {tmpl}", file=sys.stderr)
        else:
            print(f"[WARN] No DOCX template found. Place a .docx file in the 'documents/' folder.", file=sys.stderr)
        _cleanup_done = True
    return render_template('index.html')


def _cleanup_old_files():
    now = time.time()
    for f in os.listdir(OUTPUT_DIR):
        fpath = os.path.join(OUTPUT_DIR, f)
        if os.path.isfile(fpath) and now - os.path.getmtime(fpath) > 1800:
            try:
                os.remove(fpath)
            except Exception:
                pass


@app.route('/health')
def health():
    """Health check endpoint for Docker / Render."""
    tmpl = find_template()
    return jsonify({
        'status': 'ok',
        'template_found': tmpl is not None,
        'template_name': os.path.basename(tmpl) if tmpl else None
    })


@app.route('/template-status')
def template_status():
    """Check template and settings status."""
    settings = load_settings()
    tmpl = find_template()
    return jsonify({
        'settings': settings,
        'template_path': tmpl,
        'template_exists': tmpl is not None and os.path.exists(tmpl) if tmpl else False
    })


@app.route('/generate', methods=['POST'])
def generate_assignment():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data received. Please fill in all fields.', 'success': False}), 400

    student_name = data.get('studentName', '').strip()
    student_pid = data.get('studentPid', '').strip()
    student_course = data.get('studentCourse', '').strip()

    if not student_name or not student_pid or not student_course:
        return jsonify({'error': 'Please provide your name, PID, and program.', 'success': False}), 400

    file_id = uuid.uuid4().hex[:12]
    docx_path = os.path.join(OUTPUT_DIR, f'assignment_{file_id}.docx')

    try:
        tmpl = find_template()
        if not tmpl:
            return jsonify({
                'error': 'No DOCX template found. Place a .docx file in the "documents/" folder.',
                'success': False
            }), 400

        if not os.path.exists(tmpl):
            return jsonify({
                'error': f'Template file not found: {os.path.basename(tmpl)}',
                'success': False
            }), 400

        print(f"[GEN] Generating for: {student_name} | {student_pid} | {student_course}", file=sys.stderr)
        print(f"[GEN] Using template: {tmpl}", file=sys.stderr)

        modify_docx_template(student_name, student_pid, student_course, docx_path)

        if not os.path.exists(docx_path) or os.path.getsize(docx_path) == 0:
            return jsonify({
                'error': 'DOCX generation failed — output file is empty.',
                'success': False
            }), 500

        print(f"[GEN] DOCX created: {docx_path} ({os.path.getsize(docx_path)} bytes)", file=sys.stderr)

        # Convert DOCX to PDF using LibreOffice (preserves all formatting)
        pdf_path = convert_docx_to_pdf(docx_path, OUTPUT_DIR)
        pdf_ready = pdf_path is not None

        if pdf_ready:
            print(f"[GEN] PDF ready: {pdf_path} ({os.path.getsize(pdf_path)} bytes)", file=sys.stderr)
        else:
            print(f"[GEN] PDF conversion failed — DOCX was created but LibreOffice conversion failed.", file=sys.stderr)

        return jsonify({
            'success': True,
            'pdfUrl': f'/download-pdf/{file_id}' if pdf_ready else None,
            'previewUrl': f'/preview-pdf/{file_id}' if pdf_ready else None,
            'pdfReady': pdf_ready,
            'studentName': student_name,
            'studentPid': student_pid,
            'studentCourse': student_course
        })

    except FileNotFoundError as e:
        print(f"[ERR] File not found: {e}", file=sys.stderr)
        return jsonify({'error': f'Template not found: {str(e)}', 'success': False}), 400
    except Exception as e:
        print(f"[ERR] Generation failed: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return jsonify({'error': f'An error occurred: {str(e)}', 'success': False}), 500


@app.route('/preview-pdf/<file_id>')
def preview_pdf(file_id):
    """Serve the PDF inline for iframe preview."""
    pdf_path = os.path.join(OUTPUT_DIR, f'assignment_{file_id}.pdf')
    if not os.path.exists(pdf_path):
        return jsonify({'error': 'PDF not found.'}), 404

    return send_file(
        pdf_path,
        mimetype='application/pdf'
    )


@app.route('/download-pdf/<file_id>')
def download_pdf(file_id):
    pdf_path = os.path.join(OUTPUT_DIR, f'assignment_{file_id}.pdf')
    if not os.path.exists(pdf_path):
        return jsonify({'error': 'PDF not found. Please regenerate the assignment.'}), 404

    return send_file(
        pdf_path,
        as_attachment=True,
        download_name='Assignment.pdf',
        mimetype='application/pdf'
    )


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', debug=False, port=port)
